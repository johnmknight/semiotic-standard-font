"""
Update semiotic-pipeline-processes-v2.docx to reflect new OpenCV pipeline.
Rewrites Step 2 (Interior Bounds Detection) and Step 5 (Vector Tracing).
Adds new Process 1.6 section on per-icon epsilon settings.
"""
import docx
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import shutil, os

SRC = r"C:\Users\john_\dev\semiotic-standard-font\docs\semiotic-pipeline-processes-v2.docx"
DST = r"C:\Users\john_\dev\semiotic-standard-font\docs\semiotic-pipeline-processes-v3.docx"
shutil.copy(SRC, DST)

doc = Document(DST)

# Map out paragraphs we need to find and update
REPLACEMENTS = {
    # Step 2 heading trigger -> replace the body paragraph after it
    "Step 2: Interior Bounds Detection": {
        "heading": "Step 2: Cobb Border Crop (Color-Based Detection)",
        "body": (
            "The Cobb frame is detected using the COLOR source image, not grayscale brightness. "
            "All Cobb frames contain a distinctive dark-red band (R>70, G<40, B<40) regardless of the icon's interior color. "
            "This works on both light-background and dark-background icons (frozen_goods, hydroponic, etc.) where grayscale "
            "brightness approaches failed.\n\n"
            "Algorithm (scripts/step1_extract.py):\n"
            "1. Scan inward from each edge using the center-50% strip of the color image\n"
            "2. Detect rows/cols where >5% of pixels match the dark-red frame signature\n"
            "3. Find the last such row/col before 5 consecutive non-red rows (= past the red band)\n"
            "4. Add INSET=15px to clear the remaining white band and inner black border ring\n\n"
            "Outputs: work/{name}_1_source.png (extracted), work/{name}_2_bw.png (B&W), work/{name}_3_cropped.png (cropped).\n\n"
            "Preview: preview/step3_review.html shows source vs cropped for all 16 icons."
        )
    },
    "Step 5: Vector Tracing (vtracer)": {
        "heading": "Step 5: Vector Tracing (OpenCV)",
        "body": (
            "OpenCV findContours + approxPolyDP is used instead of vtracer. vtracer's polygon mode was rejected "
            "because it lost critical shape information (body width, shoulder curves) even at conservative settings.\n\n"
            "Script: scripts/trace_all.py\n\n"
            "Settings:\n"
            "  cv2.findContours(binary, RETR_CCOMP, CHAIN_APPROX_NONE)\n"
            "  cv2.approxPolyDP(contour, epsilon, closed=True)\n"
            "  fill-rule: evenodd (handles holes correctly)\n\n"
            "Per-icon epsilon values (Feb 2026):\n"
            "  epsilon=1.3 (default): allergen_warning, beverage_dispenser, contaminated, food_heating,\n"
            "                         fresh_produce, grain, hydroponic, organic_waste, potable_water,\n"
            "                         protein, utensils, water_filtration\n"
            "  epsilon=1.1 (more detail): alcohol, emergency_rations, frozen_goods, rations\n\n"
            "Output: work/{name}_4_traced.svg — single compound path, black fill, evenodd rule.\n\n"
            "Comparison tools: preview/cv_compare.html, preview/eps_compare.html, preview/flagged_compare.html"
        )
    }
}

# Walk paragraphs and apply replacements
i = 0
while i < len(doc.paragraphs):
    p = doc.paragraphs[i]
    for trigger, data in REPLACEMENTS.items():
        if trigger in p.text:
            # Update heading text
            for run in p.runs:
                if trigger in run.text:
                    run.text = run.text.replace(trigger, data["heading"])
            # Update the body paragraph(s) that follow — replace text of next non-empty para
            j = i + 1
            while j < len(doc.paragraphs):
                np_ = doc.paragraphs[j]
                if np_.text.strip():
                    # Clear existing runs and set new text
                    for run in np_.runs:
                        run.text = ""
                    if np_.runs:
                        np_.runs[0].text = data["body"]
                    else:
                        np_.add_run(data["body"])
                    break
                j += 1
    i += 1

# Add a new lessons section at the end of Process 3
# Find the last paragraph and append
# Add new section by cloning style XML from existing heading paragraphs
from lxml import etree
from copy import deepcopy

ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Find an existing Heading 2 paragraph to clone its XML structure
h2_template = None
body_template = None
for p in doc.paragraphs:
    if p.style and p.style.name == 'Heading 2' and h2_template is None:
        h2_template = deepcopy(p._element)
    if body_template is None and p.text.strip() and (p.style is None or p.style.name not in ('Heading 1','Heading 2','Heading 3')):
        body_template = deepcopy(p._element)

def make_para(template, new_text):
    el = deepcopy(template)
    for r in el.findall(f'{{{ns}}}r'):
        el.remove(r)
    r_el = etree.SubElement(el, f'{{{ns}}}r')
    t_el = etree.SubElement(r_el, f'{{{ns}}}t')
    t_el.text = new_text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return el

doc._body._body.append(make_para(h2_template,
    "3.8 Pattern: Per-Icon Epsilon Tuning for Contour Tracing"))
doc._body._body.append(make_para(body_template,
    "What works: Using a default epsilon (1.3) for most icons but dialing back to 1.1 for icons "
    "with curved or complex shapes (alcohol, emergency_rations, frozen_goods, rations). "
    "A visual review pass after batch tracing catches boxy/simplified shapes. "
    "Rule: Default epsilon=1.3. Flag icons that look angular or have lost key shape features. "
    "Re-run at epsilon=1.1. Record per-icon values in DECISIONS.md."))

doc.save(DST)
print(f"Saved: {DST}")
